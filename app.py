import csv
from collections import defaultdict
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt
import os
from datetime import datetime
from flask import Flask, render_template, request, send_file, jsonify, url_for
import urllib.parse

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = '/tmp/uploads'
app.config['DOWNLOAD_FOLDER'] = '/tmp/downloads'
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your-secret-key')

# Ensure upload and download directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DOWNLOAD_FOLDER'], exist_ok=True)

def is_section_start(row):
    return len(row) > 0 and row[0].strip('"').startswith(tuple(str(i) + '.' for i in range(1, 8)))

def create_bar_graph(speed_ranges, hours, output_path):
    print("Creating bar graph...")
    try:
        plt.figure(figsize=(5.5, 4.0), facecolor='white')
        bars = plt.bar(speed_ranges, hours, color='#4C78A8', edgecolor='black', linewidth=1.2, alpha=0.9, width=0.5)
        plt.xlabel("Engine Speed Range (r/min)", fontsize=10, labelpad=15, fontweight='bold')
        plt.ylabel("Hours", fontsize=10, labelpad=15, fontweight='bold')
        plt.xticks(rotation=45, ha='right', fontsize=8)
        plt.yticks(fontsize=8)
        plt.grid(True, axis='y', linestyle='--', alpha=0.7, color='gray')
        for bar in bars:
            yval = bar.get_height()
            plt.text(bar.get_x() + bar.get_width()/2, yval + 0.5, f'{yval}', ha='center', va='bottom', fontsize=8, fontweight='bold')
        ax = plt.gca()
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('gray')
        ax.spines['bottom'].set_color('gray')
        plt.tight_layout()
        plt.savefig(output_path, bbox_inches='tight', dpi=100)
        plt.close()
        print(f"Graph saved to {output_path}")
    except Exception as e:
        print(f"Error creating bar graph: {str(e)}")
        raise

def create_line_graph(times, hours, output_path):
    print("Creating line graph for engine oil exchange...")
    try:
        plt.figure(figsize=(3.5, 2.5), facecolor='white')
        plt.plot(times, hours, marker='o', color='#FF6F61', linewidth=2, markersize=6, markerfacecolor='#FF6F61', markeredgecolor='black', markeredgewidth=1)
        for time, hour in zip(times, hours):
            plt.text(time, hour + 50, f'{hour}', ha='center', va='bottom', fontsize=6, fontweight='bold')
        plt.xlabel("Record Number", fontsize=8, labelpad=10, fontweight='bold')
        plt.ylabel("Engine Hours", fontsize=8, labelpad=10, fontweight='bold')
        plt.xticks(times, fontsize=6)
        plt.yticks(fontsize=6)
        plt.grid(True, linestyle='--', alpha=0.7, color='gray')
        ax = plt.gca()
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.spines['left'].set_color('gray')
        ax.spines['bottom'].set_color('gray')
        plt.tight_layout()
        plt.savefig(output_path, bbox_inches='tight', dpi=100)
        plt.close()
        print(f"Line graph saved to {output_path}")
    except Exception as e:
        print(f"Error creating line graph: {str(e)}")
        raise

def remove_table_outer_borders(table):
    """Helper function to remove outer borders of a table while keeping inner borders."""
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if tblBorders is not None:
        tblBorders.clear()
    else:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    for border in ['top', 'left', 'bottom', 'right']:
        border_elem = OxmlElement(f'w:{border}')
        border_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'nil')
        tblBorders.append(border_elem)

def remove_all_table_borders(table):
    """Helper function to remove all borders (inner and outer) from a table."""
    tbl = table._element
    tblPr = tbl.tblPr
    tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if tblBorders is not None:
        tblBorders.clear()
    else:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)
    for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border_elem = OxmlElement(f'w:{border}')
        border_elem.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'nil')
        tblBorders.append(border_elem)

def process_csv_to_tables(file_path, output_dir):
    try:
        print(f"Reading CSV file: {file_path}")
        sections = defaultdict(list)
        current_section = "Metadata"
        customer_name = "Unknown"
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                csv_rows = list(reader)
                print(f"Total CSV rows read: {len(csv_rows)}")
                if not csv_rows:
                    raise ValueError("CSV file is empty")
        except UnicodeDecodeError as e:
            print(f"UnicodeDecodeError while reading CSV: {str(e)}")
            raise Exception(f"Failed to read CSV file due to encoding issue: {str(e)}")
        except csv.Error as e:
            print(f"CSV parsing error: {str(e)}")
            raise Exception(f"Error parsing CSV file: {str(e)}")
        except Exception as e:
            print(f"Unexpected error while reading CSV: {str(e)}")
            raise Exception(f"Failed to read CSV file: {str(e)}")

        for i, row in enumerate(csv_rows):
            print(f"Row {i + 1}: {row}")
            if not any(row):
                continue
            if is_section_start(row):
                current_section = row[0].strip()
                print(f"Detected section: {current_section}")
            sections[current_section].append(row)
            if current_section == "Metadata" and len(row) >= 1 and row[0].strip('"') == "Customer name":
                customer_name = row[2].strip('"') if len(row) > 2 and row[2].strip('"') else row[1].strip('"') if len(row) > 1 and row[1].strip('"') else "Unknown"
                print(f"Customer name extracted: {customer_name}")

        print("CSV file read successfully. Creating Word document...")

        today = datetime.now().strftime("%d-%m-%y")
        safe_customer_name = customer_name.replace(" ", "_").replace("/", "_").replace("\\", "_")
        output_filename = f"{safe_customer_name}_Yamaha_Diagnostics_Report_{today}.docx"
        output_file = os.path.join(output_dir, output_filename)
        print(f"Output file will be saved as: {output_file}")

        # Extract Total Engine Hours
        total_engine_hours = "(empty)"
        if "1. Engine operating hours according to engine speed" in sections:
            print("Section 1 data:", sections["1. Engine operating hours according to engine speed"])
            for row in sections["1. Engine operating hours according to engine speed"]:
                if len(row) >= 1 and row[0].strip('"') == "Total operating hours":
                    total_engine_hours = row[2].strip('"').strip() if len(row) > 2 and row[2].strip('"') else row[1].strip('"').strip() if len(row) > 1 and row[1].strip('"') else "(empty)"
                    break
        print(f"Total Engine Hours extracted: {total_engine_hours}")

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        print("Available table styles:")
        for style in doc.styles:
            if style.type == 2:
                print(style.name)

        logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
        if os.path.exists(logo_path):
            header = doc.sections[0].header
            logo_paragraph = header.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_paragraph.add_run()
            run.add_picture(logo_path, width=Inches(3))
            doc.sections[0].header_distance = Inches(0.1)
        else:
            print(f"Error: logo.png not found at {logo_path}. Skipping logo.")

        # Metadata section (center-aligned, no borders, larger text, no dotted line)
        if "Metadata" in sections:
            print("Processing Metadata section...")
            metadata_entries = []
            metadata_keys = [
                "YAMAHA DIAGNOSTIC SYSTEM", "Save date & time", "Customer name", 
                "Dealer name", "Number of engines", "Comment", "Model name", 
                "Engine serial number (PID number)", "ECM number"
            ]
            comment_value = "(empty)"
            i = 0
            while i < len(sections["Metadata"]):
                row = sections["Metadata"][i]
                if len(row) >= 1 and row[0].strip('"') in metadata_keys:
                    field = row[0].strip('"')
                    if field == "Comment":
                        if i + 1 < len(sections["Metadata"]) and not is_section_start(sections["Metadata"][i + 1]):
                            next_row = sections["Metadata"][i + 1]
                            comment_value = next_row[0].strip('"') if len(next_row) > 0 and next_row[0].strip('"') else "(empty)"
                            i += 1
                    else:
                        value = row[2].strip('"') if len(row) > 2 and row[2].strip('"') else row[1].strip('"') if len(row) > 1 and row[1].strip('"') != field else "(empty)"
                        display_field = "Service Date" if field == "Save date & time" else field
                        if field == "Dealer name":
                            value = "Northside Marine"
                        if value != "(empty)":
                            metadata_entries.append((display_field, value))
                i += 1
            if comment_value != "(empty)":
                metadata_entries.append(("Comment", comment_value))
            if total_engine_hours != "(empty)":
                metadata_entries.append(("Total Engine Hours", total_engine_hours))

            if metadata_entries:
                print(f"Creating Metadata table with {len(metadata_entries)} entries")
                num_rows = (len(metadata_entries) + 1) // 2
                table = doc.add_table(rows=num_rows, cols=2)
                table.style = 'Table Grid'
                remove_all_table_borders(table)  # Remove all borders (no dotted line)
                table.autofit = True
                for idx, (field, value) in enumerate(metadata_entries):
                    row_idx = idx // 2
                    col_idx = idx % 2
                    cell = table.cell(row_idx, col_idx)
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_field = paragraph.add_run(f"{field}: ")
                    run_field.bold = True
                    run_field.font.size = Pt(10)
                    run_value = paragraph.add_run(value)
                    run_value.bold = False
                    run_value.font.size = Pt(10)
                    paragraph.space_after = Pt(2)

                # Removed the single line space after the Metadata section
                # doc.add_paragraph()

        # Main content: Two columns (graphs on left, tables on right)
        main_table = doc.add_table(rows=2, cols=2)
        main_table.style = 'Table Grid'
        remove_all_table_borders(main_table)  # Remove all borders to eliminate lines between graphs and tables
        main_table.autofit = True

        # Top Left: Engine Operating Hours bar graph
        cell_top_left = main_table.cell(0, 0)
        heading_paragraph = cell_top_left.add_paragraph()
        heading_run = heading_paragraph.add_run("Engine Operating Hours According to Engine Speed")
        heading_run.bold = True
        heading_run.font.size = Pt(10)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        graph_paragraph = cell_top_left.add_paragraph()
        graph_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        speed_ranges = []
        hours = []
        if "1. Engine operating hours according to engine speed" in sections:
            for row in sections["1. Engine operating hours according to engine speed"]:
                if len(row) >= 2 and "r/min" in row[0]:
                    speed_range = row[0].strip('"')
                    try:
                        hour = float(row[2].strip('"') if len(row) > 2 and row[2].strip('"') else "0")
                        if hour > 0:
                            speed_ranges.append(speed_range)
                            hours.append(hour)
                    except ValueError as e:
                        print(f"Warning: Could not convert '{row[2]}' to float in Section 1: {str(e)}")
                        continue

        if speed_ranges and hours:
            graph_path = os.path.join(output_dir, "engine_hours_graph.png")
            create_bar_graph(speed_ranges, hours, graph_path)
            run = graph_paragraph.add_run()
            run.add_picture(graph_path, width=Inches(3.5), height=Inches(4.0))
            os.remove(graph_path)
        else:
            cell_top_left.add_paragraph("No significant operating hours to display.")

        # Bottom Left: Engine Oil Exchange line graph
        cell_bottom_left = main_table.cell(1, 0)
        heading_paragraph = cell_bottom_left.add_paragraph()
        heading_run = heading_paragraph.add_run("Record of Engine Oil Exchange")
        heading_run.bold = True
        heading_run.font.size = Pt(10)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        times = []
        hours = []
        if "2. Record of engine oil exchange" in sections:
            for row in sections["2. Record of engine oil exchange"]:
                if len(row) >= 2 and row[0].strip('"') != "Time" and row[0].strip('"'):
                    time = row[0].strip('"')
                    data = row[2].strip('"') if len(row) > 2 and row[2].strip('"') else "(empty)"
                    if data != "(empty)":
                        try:
                            times.append(int(time))
                            hours.append(float(data))
                        except ValueError as e:
                            print(f"Warning: Could not convert time '{time}' or hours '{data}' to number: {str(e)}")
                            continue

        if times and hours:
            line_graph_path = os.path.join(output_dir, "engine_oil_exchange_graph.png")
            create_line_graph(times, hours, line_graph_path)
            graph_paragraph = cell_bottom_left.add_paragraph()
            graph_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = graph_paragraph.add_run()
            run.add_picture(line_graph_path, width=Inches(3.5), height=Inches(2.5))
            os.remove(line_graph_path)
        else:
            no_data_paragraph = cell_bottom_left.add_paragraph("No engine oil exchange records to display.")
            no_data_paragraph.runs[0].font.size = Pt(8)
            no_data_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Right Column: Merge the two cells and add Engine Record and Engine Monitor tables
        cell_top_right = main_table.cell(0, 1)
        cell_bottom_right = main_table.cell(1, 1)
        cell_top_right.merge(cell_bottom_right)  # Merge the two cells into one
        cell_right = cell_top_right  # Use the merged cell

        # Engine Record Section
        heading_paragraph = cell_right.add_paragraph()
        heading_run = heading_paragraph.add_run("Engine Record")
        heading_run.bold = True
        heading_run.font.size = Pt(10)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add space between header and table
        cell_right.add_paragraph()

        has_data = False
        table_data = []
        if "6. Engine record" in sections:
            for row in sections["6. Engine record"]:
                if len(row) >= 1 and row[0].strip('"') != "Data item" and row[0].strip('"') != "6. Engine record" and row[0].strip('"'):
                    item = row[0].strip('"')
                    value = row[2].strip('"') if len(row) > 2 and row[2].strip('"') else row[1].strip('"') if len(row) > 1 and row[1].strip('"') else "(empty)"
                    if item != "(empty)" and value != "(empty)" and value.strip():  # Only include if value exists
                        table_data.append((item, value))
                        has_data = True

        if has_data:
            sub_table = cell_right.add_table(rows=1, cols=2)
            sub_table.style = 'Table Grid'
            remove_table_outer_borders(sub_table)
            sub_table.autofit = True
            hdr_cells = sub_table.rows[0].cells
            hdr_cells[0].text = "Data Item"
            hdr_cells[0].paragraphs[0].runs[0].bold = True
            hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(8)
            hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[1].text = "Value"
            hdr_cells[1].paragraphs[0].runs[0].bold = True
            hdr_cells[1].paragraphs[0].runs[0].font.size = Pt(8)
            hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for item, value in table_data:
                row_cells = sub_table.add_row().cells
                row_cells[0].text = item
                row_cells[0].paragraphs[0].runs[0].font.size = Pt(8)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[1].text = value
                row_cells[1].paragraphs[0].runs[0].font.size = Pt(8)
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            no_data_paragraph = cell_right.add_paragraph("No engine records to display.")
            no_data_paragraph.runs[0].font.size = Pt(8)
            no_data_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add space between Engine Record and Engine Monitor
        cell_right.add_paragraph()

        # Engine Monitor Section
        heading_paragraph = cell_right.add_paragraph()
        heading_run = heading_paragraph.add_run("Engine Monitor")
        heading_run.bold = True
        heading_run.font.size = Pt(10)
        heading_run.font.color.rgb = RGBColor(0, 0, 0)
        heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add space between header and table
        cell_right.add_paragraph()

        has_data = False
        table_data = []
        stop_adding = False
        if "4. Engine monitor" in sections:
            for row in sections["4. Engine monitor"]:
                if len(row) >= 5 and row[0].strip('"') and not row[0].startswith("Monitor item") and not row[0].startswith("4. Engine monitor"):
                    item = row[0].strip('"')
                    if item.lower() == "engine shut off switch":
                        stop_adding = True
                        break
                    if stop_adding:
                        continue
                    # Skip items containing A/D(CH1), A/D(CH2), A/D(CH3) in any case
                    if any(ch in item.lower() for ch in ["a/d(ch1)", "a/d(ch2)", "a/d(ch3)"]):
                        continue
                    value = row[4].strip('"') if len(row) > 4 else row[2].strip('"')
                    if value != "(empty)" and value.strip():  # Only include if value exists
                        table_data.append((item, value))
                        has_data = True

        if has_data:
            sub_table = cell_right.add_table(rows=1, cols=2)
            sub_table.style = 'Table Grid'
            remove_table_outer_borders(sub_table)
            sub_table.autofit = True
            hdr_cells = sub_table.rows[0].cells
            hdr_cells[0].text = "Monitor Item"
            hdr_cells[0].paragraphs[0].runs[0].bold = True
            hdr_cells[0].paragraphs[0].runs[0].font.size = Pt(8)
            hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[1].text = "Value"
            hdr_cells[1].paragraphs[0].runs[0].bold = True
            hdr_cells[1].paragraphs[0].runs[0].font.size = Pt(8)
            hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for item, value in table_data:
                row_cells = sub_table.add_row().cells
                row_cells[0].text = item
                row_cells[0].paragraphs[0].runs[0].font.size = Pt(8)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[1].text = value
                row_cells[1].paragraphs[0].runs[0].font.size = Pt(8)
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            no_data_paragraph = cell_right.add_paragraph("No engine monitor data to display.")
            no_data_paragraph.runs[0].font.size = Pt(8)
            no_data_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Diagnosis section
        if "3. Diagnosis" in sections:
            print("Processing Diagnosis section...")
            heading_paragraph = doc.add_paragraph()
            heading_run = heading_paragraph.add_run("Diagnosis")
            heading_run.bold = True
            heading_run.font.size = Pt(10)
            heading_run.font.color.rgb = RGBColor(0, 0, 0)
            heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            has_data = False
            table_data = []
            for row in sections["3. Diagnosis"]:
                if len(row) >= 3 and row[2].strip('"').isdigit():
                    item = row[0].strip('"')
                    status = row[1].strip('"')
                    code = row[2].strip('"')
                    if status != "(empty)" and code != "(empty)" and status.strip() and code.strip():  # Only include if both values exist
                        table_data.append((item, status, code))
                        has_data = True

            if has_data:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                remove_table_outer_borders(table)
                table.autofit = True
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Item"
                hdr_cells[0].paragraphs[0].runs[0].bold = True
                hdr_cells[1].text = "Status"
                hdr_cells[1].paragraphs[0].runs[0].bold = True
                hdr_cells[2].text = "Code"
                hdr_cells[2].paragraphs[0].runs[0].bold = True
                for item, status, code in table_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = item
                    row_cells[1].text = status
                    row_cells[2].text = code
            else:
                doc.add_paragraph("No diagnosis records to display.")
            doc.add_paragraph()

        print(f"Saving Word document to {output_file}")
        doc.save(output_file)
        return output_file

    except Exception as e:
        print(f"An error occurred in process_csv_to_tables: {str(e)}")
        raise

# Flask routes
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    max_file_size = 5 * 1024 * 1024
    if int(request.headers.get('Content-Length', 0)) > max_file_size:
        return jsonify({'success': False, 'message': 'File too large. Maximum size is 5MB.'}), 400

    if 'file' not in request.files:
        return jsonify({'success': False, 'message': 'No file part in the request.'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No file selected.'}), 400

    if file and file.filename.endswith('.csv'):
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], 'uploaded.csv')
        print(f"Saving uploaded file to: {upload_path}")
        file.save(upload_path)

        try:
            output_file = process_csv_to_tables(upload_path, app.config['DOWNLOAD_FOLDER'])
            output_filename = os.path.basename(output_file)

            print(f"Cleaning up old files in {app.config['DOWNLOAD_FOLDER']}")
            for f in os.listdir(app.config['DOWNLOAD_FOLDER']):
                f_path = os.path.join(app.config['DOWNLOAD_FOLDER'], f)
                if f_path != output_file:
                    try:
                        if os.path.isfile(f_path):
                            os.remove(f_path)
                            print(f"Removed old file: {f_path}")
                    except Exception as e:
                        print(f"Error removing old file {f_path}: {str(e)}")

            print(f"Removing uploaded file: {upload_path}")
            os.remove(upload_path)

            download_url = url_for('download_file', filename=urllib.parse.quote(output_filename))
            print(f"Generated download URL: {download_url}")
            return jsonify({'success': True, 'download_url': download_url})
        except Exception as e:
            print(f"Error in /process endpoint: {str(e)}")
            return jsonify({'success': False, 'message': f"Error processing file: {str(e)}"}), 500
    else:
        return jsonify({'success': False, 'message': 'Please upload a valid CSV file.'}), 400

@app.route('/download/<filename>')
def download_file(filename):
    file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
    print(f"Serving file for download: {file_path}")
    return send_file(file_path, as_attachment=True)

@app.route('/logo')
def serve_logo():
    logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
    if os.path.exists(logo_path):
        return send_file(logo_path, mimetype='image/png')
    else:
        return "Logo not found", 404

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=int(os.environ.get('PORT', 5000)))
